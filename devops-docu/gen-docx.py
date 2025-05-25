import os
from docx import Document

FOLDER = "devops-docu"  # <-- Replace with your folder name
TEMPLATE = "template.docx"
OUTPUT = "output.docx"

doc = Document(TEMPLATE)
doc.add_heading(f'Contents of {FOLDER}', 0)

for filename in sorted(os.listdir(FOLDER)):
    filepath = os.path.join(FOLDER, filename)
    if os.path.isfile(filepath):
        doc.add_heading(filename, level=1)
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        doc.add_paragraph(content)

doc.save(OUTPUT)